import json
import base64
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO


def add_table_borders(table):
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)


def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    '''
    Business: Creates Word document with welding control data in standard format
    Args: event - dict with httpMethod, body containing formData and connections
          context - object with request_id attribute
    Returns: HTTP response with base64-encoded .docx file
    '''
    method: str = event.get('httpMethod', 'POST')
    
    if method == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Max-Age': '86400'
            },
            'body': ''
        }
    
    if method != 'POST':
        return {
            'statusCode': 405,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'isBase64Encoded': False,
            'body': json.dumps({'error': 'Method not allowed'})
        }
    
    body_data = json.loads(event.get('body', '{}'))
    form_data = body_data.get('formData', {})
    connections = body_data.get('connections', [])
    
    doc = Document()
    
    p = doc.add_paragraph()
    run = p.add_run(f"Наименование лаборатории:             {form_data.get('labName', 'ЛККСС')}")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Свидетельство об аттестации:          {form_data.get('certificate', '')}")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Нормативный документ:                   {form_data.get('normDoc', 'СТО Газпром 15-1.3-004-2023')}")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Номер ОТК НК:                                  {form_data.get('otkNumber', '')}")
    run.font.size = Pt(11)
    
    table = doc.add_table(rows=4, cols=2)
    add_table_borders(table)
    
    table.cell(0, 0).text = 'Средства контроля'
    table.cell(0, 0).merge(table.cell(3, 0))
    
    table.cell(0, 1).text = 'Источник ионизирующего излучения'
    table.cell(1, 1).text = 'Тип детектора'
    table.cell(2, 1).text = 'Тип защитного экрана'
    table.cell(3, 1).text = 'Тип усиливающего экрана'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run(f"Заключение по ВИК: № {form_data.get('vikNumber', '')}     от  {form_data.get('vikDate', '')} г.")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Наименование объекта: {form_data.get('objectName', '')}")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Уровень качества:                           {form_data.get('qualityLevel', 'В')}")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Объём контроля:                               {form_data.get('controlVolume', '100')}")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Название трассы: {form_data.get('routeName', '')}")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Участок труб-да, километраж:       {form_data.get('pipelineSection', '-')}")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Наименование орг. подрядчика:    {form_data.get('contractor', '')}")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Наименование орг. заказчика:       {form_data.get('customer', '')}")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Шифр бригады или клеймо сварщика: {form_data.get('welderCode', '')}")
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"ЗАКЛЮЧЕНИЕ № {form_data.get('conclusionNumber', '')}")
    run.bold = True
    run.font.size = Pt(12)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"от «{form_data.get('conclusionDate', '')}»")
    run.font.size = Pt(11)
    
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('по результатам контроля качества сварных соединений радиационным неразрушающим методом (РК)')
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('РЕЗУЛЬТАТЫ КОНТРОЛЯ:')
    run.bold = True
    run.font.size = Pt(11)
    
    if connections:
        table = doc.add_table(rows=1, cols=9)
        add_table_borders(table)
        
        hdr_cells = table.rows[0].cells
        headers = [
            'Номер сварного соединения\nпо журналу сварки',
            'Диаметр и\nтолщина (номинальная/\nрадиационная) стенки трубы, мм',
            'Шифр бригады или клеймо сварщика',
            'Номер участка контроля,\n(координаты мерного пояса),\nмм',
            'Чувстви-тельность контроля,\nмм',
            'Координаты выявленных дефектов по периметру сварного шва, мм',
            'ОПИСАНИЕ ВЫЯВЛЕНЫХ ДЕФЕКТОВ',
            'ЗАКЛЮЧЕНИЕ (годен, ремонт, вырезать)',
            'Примеча-ние'
        ]
        
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.bold = True
        
        for conn in connections:
            row_cells = table.add_row().cells
            row_cells[0].text = str(conn.get('number', ''))
            row_cells[1].text = str(conn.get('diameter', ''))
            row_cells[2].text = str(conn.get('welderCode', ''))
            row_cells[3].text = str(conn.get('section', ''))
            row_cells[4].text = str(conn.get('sensitivity', ''))
            row_cells[5].text = str(conn.get('coordinates', ''))
            row_cells[6].text = str(conn.get('defects', ''))
            row_cells[7].text = str(conn.get('conclusion', ''))
            row_cells[8].text = str(conn.get('note', ''))
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run(f"Заключение о качестве сварного соединения:   {form_data.get('finalConclusion', 'годен')}")
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run('                                                                              (годен, ремонт, вырезать)')
    run.font.size = Pt(9)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    info_table = doc.add_table(rows=2, cols=5)
    add_table_borders(info_table)
    
    info_table.cell(0, 0).text = 'Контроль провел'
    info_table.cell(0, 1).text = f"Ф.И.О.                {form_data.get('controllerName', '')}"
    info_table.cell(0, 2).text = f"Уровень квал.:   {form_data.get('controllerLevel', '')}  ,     № удост.: {form_data.get('controllerCertificate', '')}"
    info_table.cell(0, 3).text = 'Подпись:'
    info_table.cell(0, 4).text = f"Дата:      {form_data.get('controlDate', '')}"
    
    info_table.cell(1, 0).text = 'Заключение выдал'
    info_table.cell(1, 1).text = f"Ф.И.О.                {form_data.get('issuerName', form_data.get('controllerName', ''))}"
    info_table.cell(1, 2).text = f"Уровень квал.:   {form_data.get('issuerLevel', form_data.get('controllerLevel', ''))}  ,     № удост.: {form_data.get('issuerCertificate', form_data.get('controllerCertificate', ''))}"
    info_table.cell(1, 3).text = 'Подпись:'
    info_table.cell(1, 4).text = f"Дата:      {form_data.get('issueDate', form_data.get('controlDate', ''))}"
    
    for row in info_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    
    result_base64 = base64.b64encode(output_stream.read()).decode('utf-8')
    
    return {
        'statusCode': 200,
        'headers': {
            'Content-Type': 'application/json',
            'Access-Control-Allow-Origin': '*'
        },
        'isBase64Encoded': False,
        'body': json.dumps({
            'file': result_base64,
            'filename': f"zaklyuchenie_{form_data.get('conclusionNumber', 'rk')}.docx"
        })
    }
